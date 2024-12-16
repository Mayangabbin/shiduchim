import pandas as pd
import os
import hashlib
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
import pickle
from docx import Document  # Import the python-docx library to create Word files
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# Authenticate and create the Drive service
def authenticate_drive():
    SCOPES = ['https://www.googleapis.com/auth/drive.file']  # Permission to upload files
    creds = None
    if os.path.exists('token.pickle'):
        with open('token.pickle', 'rb') as token:
            creds = pickle.load(token)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        with open('token.pickle', 'wb') as token:
            pickle.dump(creds, token)
    return build('drive', 'v3', credentials=creds)

# Upload a file to Google Drive
def upload_to_drive(file_path, drive_service, parent_folder_id=None):
    file_metadata = {
        'name': os.path.basename(file_path),
        'mimeType': 'application/vnd.google-apps.document'  # Set mimeType for Google Docs
    }
    if parent_folder_id:
        file_metadata['parents'] = [parent_folder_id]
    
    media = MediaFileUpload(file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')  # Upload as DOCX
    uploaded_file = drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    print(f"Uploaded file {file_path} to Google Drive with ID: {uploaded_file.get('id')}")
    return uploaded_file.get('id')

# Download an image from Google Drive
def download_image_from_drive(file_id, drive_service, destination_path):
    request = drive_service.files().get_media(fileId=file_id)
    fh = io.FileIO(destination_path, 'wb')
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while done is False:
        status, done = downloader.next_chunk()
    print(f"Downloaded image from Google Drive with ID {file_id} to {destination_path}")

# Insert image into the Word document
def add_image_to_doc(doc, image_path):
    # Add image to document (adjust size as needed)
    doc.add_picture(image_path, width=Inches(3.0))
    doc.add_paragraph("\n")  # Add space after the image
    print(f"Image {image_path} added to document.")

# File paths
file_path = "responses.csv" # Input CSV file path
updated_file_path = 'responses_with_ids.csv'  # Updated CSV with Response IDs
processed_ids_file = 'processed_ids.txt'  # File to track processed Response IDs

# Load the responses CSV file
responses = pd.read_csv(file_path)

# Define a function to generate a unique Response ID
def generate_response_id(row):
    unique_string = f"generated_{row['שם פרטי']}{row['שם משפחה']}{row['חותמת זמן']}"  # Example fields
    return hashlib.md5(unique_string.encode()).hexdigest()  # Create a unique hash

# Check if Response ID column already exists; if not, add it
if 'Response ID' not in responses.columns:
    responses['Response ID'] = responses.apply(generate_response_id, axis=1)
    # Save the updated responses with Response IDs
    responses.to_csv(updated_file_path, index=False)
    print(f"Updated CSV saved to {updated_file_path} with Response IDs.")

# Load processed IDs
if os.path.exists(processed_ids_file):
    with open(processed_ids_file, 'r') as file:
        processed_ids = set(file.read().splitlines())
else:
    processed_ids = set()

# Define a function to generate a personalized text
def generate_text(row):
    return rf"""
    {row['שם פרטי']} {row['שם משפחה']}
    מספר הטלפון שלך: {row['מספר הטלפון שלך']}
    מספרים לבירורים: {row['מספרים לבירורים']}
    תאריך לידה: {row['תאריך לידה']}
    גובה: {row['גובה']}
    
    גדלתי ב{row['איפה גדלת?']}
    כרגע גר ב{row['מקום מגורים נוכחי']}
    למדתי בתיכון ב{row['איפה למדת בתיכון?']}
    בהמשך למדתי ב{row['באיזו ישיבה או מכינה למדת? כמה שנים?']}
    
    תחנות בחיים: {row['תחנות בחיים']}
    עיסוק נוכחי: {row['עיסוק נוכחי']}
    
    על המשפחה:
    {row['ספר בקווים כלליים על המשפחה שלך']}

    אופי:
    תכונות שמאפיינות אותי:
    {row['4 תכונות לפחות שמאפיינות אותך']}
    מבחינה חברתית:
    {row['מי אני מבחינה חברתית וסוג שיח..']}
    
    אני {row['אתה יותר ספונטני או מתוכנן?']}, {row['שקט או דומיננטי?']},
    סגנון: {row['אתה בסגנון עירוני או ישובי?']}
    
    בזמני הפנוי- {row['מה משמח אותך? מה אתה אוהב לעשות? מה אתה עושה בזמן פנוי או חופשות?']}
    
    שאיפות אידיאלים ותכנונים לעתיד:
    {row['אידיאלים, שליחות מדבר אלייך?']}
    {row['שאיפות ותכנונים לעתיד']}
    
    רמה וסגנון תורני:
    {row['ספר לנו על הרמה והסגנון התורני שלך']}
    רוצה לבנות-
    {row['איזה סגנון בית אתה רוצה לבנות?']}    

    מה אתה מחפש:
    {row['תכונות שמהותיות וקריטיות לך בבחורה']}
    
    {row['וקצת יותר בהרחבה..']}

    הכי חשוב לי-
    {row['לסיכום אם היית צריך לציין בנקודות מהם שלושת הדברים שהכי חשובים לך בבחורה, מה הם?']}
    
    מבחינה תורנית-
    {row['מה אתה מחפש אצל בחורה מבחינה תורנית?']}

    הערות נוספות-
    {row['הערות נוספות או מידע שיכול להועיל?']}

    האם חשובה לך העדה?: {row['האם חשובה לך העדה?']}
    האם יש משהו שאפסול עליו מבחינה חיצונית?: {row['האם יש משהו שאפסול עליו מבחינה חיצונית?']}
    
    טווח גילאים רלוונטי: {row['טווח גילאים רלוונטי']}
    אם יש מסקנות מפגישות קודמות מה נכון וטוב לי, ומה פחות, כתוב לנו. זה יועיל בלדייק אותנו.:
    {row['אם יש מסקנות מפגישות קודמות מה נכון וטוב לי, ומה פחות,  כתוב לנו. זה יועיל בלדייק אותנו.']}
    
    משהו נוסף?:
    {row['משהו נוסף שתרצה להוסיף בהקשר של מה אתה מחפש?']}

    מוכן לשלוח תמונה?
    {row['האם אתה מוכן לשלוח תמונה שלך לבחורה במקרה והיא מעוניינת בכך? ']}
    
    """ 

# Function to add a paragraph with RTL alignment
def add_rtl_paragraph(doc, text):
    # Add paragraph
    para = doc.add_paragraph(text)
    
    # Set alignment to right
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Set the RTL direction (this requires setting the 'right-to-left' property)
    run = para.add_run()
    run._r.set('rtl', '1')  # This sets the text direction to right-to-left
    
    return para

# Authenticate Google Drive service
drive_service = authenticate_drive()

# Process new responses
new_ids = []
for index, row in responses.iterrows():
    response_id = row['Response ID']
    if response_id not in processed_ids:  # Process only unprocessed responses
        summary = generate_text(row)
        
        # Create a .docx file using python-docx
        doc = Document()

        # Add the content as RTL paragraphs
        add_rtl_paragraph(doc, summary)
        
        # Check if there's an image to insert (e.g., if there's an image file ID in the row)
        if 'Image ID' in row and pd.notna(row['Image ID']):
            image_id = row['Image ID']
            image_path = f"image_{response_id}.jpg"
            download_image_from_drive(image_id, drive_service, image_path)
            add_image_to_doc(doc, image_path)
            os.remove(image_path)  # Optionally remove the image after adding it to the document
        
        # Save the document locally as a .docx file
        output_file = f"generated_{row['שם פרטי']}_{row['שם משפחה']}_{row['מספר הטלפון שלך']}.docx"
        doc.save(output_file)

        # Upload the file to Google Drive as a Google Doc
        file_id = upload_to_drive(output_file, drive_service)
        print(f"Uploaded profile for {row['שם פרטי']} {row['שם משפחה']} to Google Drive with ID: {file_id}")
        
        # Optionally delete the local file after uploading
        os.remove(output_file)
        print(f"Profile for {row['שם פרטי']}{row['שם משפחה']}{row['מספר הטלפון שלך']} saved to {output_file}")
        
        new_ids.append(response_id)

# Update the processed IDs file
if new_ids:
    with open(processed_ids_file, 'a') as file:
        file.write('\n'.join(new_ids) + '\n')

import pandas as pd
import os
import hashlib
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
import pickle
from docx import Document  # Import the python-docx library to create Word files
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import google.generativeai as genai
from google.oauth2 import service_account
from googleapiclient.discovery import build
import re
from PIL import Image
import pyheif


# Authenticate and create the Drive service
def authenticate_drive():
    # Path to your service account JSON key
    SERVICE_ACCOUNT_FILE = "./credentials.json"

    SCOPES = ['https://www.googleapis.com/auth/drive']

    # Authenticate using service account
    creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)

    return build('drive', 'v3', credentials=creds)

# Function to download the Google Sheets file as CSV
def download_csv_from_drive(file_id, drive_service, destination_path):
    # Export the Google Sheet as CSV
    request = drive_service.files().export_media(fileId=file_id, mimeType='text/csv')
    fh = io.FileIO(destination_path, 'wb')
    downloader = MediaIoBaseDownload(fh, request)
    
    done = False
    while done is False:
        status, done = downloader.next_chunk()
    print(f"Downloaded CSV from Google Drive with File ID {file_id} to {destination_path}")

# Upload a file to Google Drive
def upload_to_drive(file_path, drive_service, parent_folder_id='1uoXnNThQwdB-C6IeKV7NKm-kX-22B8t0'):
    file_metadata = {
        'name': os.path.basename(file_path),
        'mimeType': 'application/vnd.google-apps.document'  # Set mimeType for Google Docs
    }
    if parent_folder_id:
        file_metadata['parents'] = [parent_folder_id]
    
    media = MediaFileUpload(file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')  # Upload as DOCX
    uploaded_file = drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    print(f"Uploaded file {file_path} to Google Drive with ID: {uploaded_file.get('id')}")
    return uploaded_file.get('id')

# Download an image from Google Drive
def download_image_from_drive(file_id, drive_service, destination_path):
    request = drive_service.files().get_media(fileId=file_id)
    fh = io.FileIO(destination_path, 'wb')
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while done is False:
        status, done = downloader.next_chunk()

    try:
        # Try opening the image with Pillow
        try:
            img = Image.open(destination_path)
            print(f"Image format detected by Pillow: {img.format}")
        except Exception as e:
            print(f"Error opening with Pillow: {e}")
            # If Pillow fails, attempt to open with pyheif
            print("Attempting to open with pyheif (HEIF format)")
            heif_file = pyheif.read(destination_path)
            img = Image.frombytes(
                heif_file.mode, heif_file.size, heif_file.data, "raw", heif_file.mode, heif_file.stride, heif_file.size[1]
            )
            print("Opened HEIF image successfully.")

        # Ensure image is converted to a standard format (JPEG)
        if img.format is None:  # This means it's still in raw format
            print("Image format is None, converting to RGB and saving as JPEG")
            img = img.convert("RGB")  # Convert to RGB to save as JPEG
            img.save(destination_path, "JPEG")
            print(f"Converted image saved as {destination_path}")
        else:
            print(f"Image format is {img.format}, no conversion needed.")

    except Exception as e:
        print(f"Error processing the image: {e}")

    print(f"Downloaded image from Google Drive with ID {file_id} to {destination_path}")

# Insert image into the Word document
def add_image_to_doc(doc, image_path):
    # Add image to document (adjust size as needed)
    doc.add_picture(image_path, width=Inches(1.5))
    doc.add_paragraph("\n")  # Add space after the image
    print(f"Image {image_path} added to document.")

def add_images_side_by_side(doc, image_paths):
    # Create a new paragraph for the images
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align the paragraph
    
    for image_path in image_paths:
        run = paragraph.add_run()
        print (f"addind img {image_path} to file")
        run.add_picture(image_path, height=Cm(5))  # Add image to the run
        run.add_text(" ")  # Add a space between images (optional)

# Extract file IDs from Google Drive URLs
def extract_file_ids(image_urls):
    urls = image_urls.split(', ')
    file_ids = [url.split('id=')[-1] for url in urls if 'id=' in url]
    return file_ids

# File paths
responses_path = "responses1.csv" # Input CSV file path
updated_file_path = 'responses_with_ids.csv'  # Updated CSV with Response IDs
processed_ids_file = 'processed_ids.txt'  # File to track processed Response IDs

gemini_api_key = "AIzaSyD6y0TKHwbHieDps-Kjamtb-2cTYxvQBHk"

# Example usage:
file_id = "1YYMjOniotZhq_32kCHKxXSObYOO3UHB0KO-FFvuD5hY"  # Replace with the actual File ID of your Google Sheet

# Authenticate Google Drive service
drive_service = authenticate_drive()

# Download the CSV
download_csv_from_drive(file_id, drive_service, responses_path)

# Load the responses CSV file
responses = pd.read_csv(responses_path)

# Define a function to generate a unique Response ID
def generate_response_id(row):
    unique_string = f"generated_{row['שם פרטי']}{row['שם משפחה']}{row['חותמת זמן']}"  # Example fields
    return hashlib.md5(unique_string.encode()).hexdigest()  # Create a unique hash

# Check if Response ID column already exists; if not, add it
if 'Response ID' not in responses.columns:
    responses['Response ID'] = responses.apply(generate_response_id, axis=1)
    # Save the updated responses with Response IDs
    responses.to_csv(updated_file_path, index=False)
    print(f"Updated CSV saved to {updated_file_path} with Response IDs.")

# Load processed IDs
if os.path.exists(processed_ids_file):
    with open(processed_ids_file, 'r') as file:
        processed_ids = set(file.read().splitlines())
else:
    processed_ids = set()
### trying to use ai
def improve_text_with_gemini(api_key, text):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-1.5-flash")
    role_content = """
אתה כותב טקסטים בעברית בצורה מקצועית, תוך הקפדה על שמירת המבנה המקורי של הכותרות והפסקאות. הכותרות מוגדרות מראש, ואין לשנות אותן. יש לשפר רק את התוכן מתחת לכל כותרת תוך שמירה על כבוד לסגנון המקורי.

כותרות המשנה הן:
1. :שם
2. :תחנות בחיים
3. :על המשפחה
4. :אופי
5. :שאיפות, אידיאלים ותכנונים לעתיד
6. :רמה וסגנון תורני
7. :מה אתה מחפש
8. :הכי חשוב לי
9. :בהרחבה
10. :מבחינה תורנית
11. :אפסול חיצונית על
12. :טווח גילאים
13. :האם מוכן לשלוח תמונה

עליך:
- להשאיר את הכותרות כפי שהן, ללא שינוי.
- להימנע מתוספות שאינן קיימות בטקסט המקורי, ולשנות אותו כמה שפחות, לשנות אותו אך ורק כשיש בעיה תחבירית
- כדי לסמן כותרות תתחיל שורה במקף ואז רווח בודד ולאחר הכותרת נקודותיים ואז ירידת שורה 
- לא לשנות את הטקסט!!!! רק לתקן טעויות תחביר
- כמה שפחות לשנות
- כמובן שאם יש nan אז לא לכלול אותו
בבקשה שפר את הטקסט הבא לפי ההנחיות.

"""
    chat = model.start_chat(
        history=[
            {"role": "user", "parts": role_content},
            {"role": "model", "parts": "Great to meet you. What would you like to know?"},
        ]
    )

    response = chat.send_message(text)
    return(response.text)

# Define a function to generate a personalized text
def generate_text(row):
    template_text = rf"""
    {row['שם פרטי']} {row['שם משפחה']}
    מספר הטלפון שלך: {row['מספר הטלפון שלך']}
    מספרים לבירורים: {row['מספרים לבירורים']}
    תאריך לידה: {row['תאריך לידה']}
    גובה: {row['גובה']}
    
    גדלתי ב{row['איפה גדלת?']}
    כרגע גר ב{row['מקום מגורים נוכחי']}
    למדתי בתיכון ב{row['איפה למדת בתיכון?']}
    בהמשך למדתי ב{row['באיזו ישיבה או מכינה למדת? כמה שנים?']}
    
    תחנות בחיים: {row['תחנות בחיים']}, כיום {row['עיסוק נוכחי']}
    
    על המשפחה:
    {row['ספר בקווים כלליים על המשפחה שלך']}

    אופי:
    תכונות שמאפיינות אותי:
    {row['4 תכונות לפחות שמאפיינות אותך']}
    מבחינה חברתית:
    {row['מי אני מבחינה חברתית וסוג שיח..']}
    
    אני {row['אתה יותר ספונטני או מתוכנן?']}, {row['שקט או דומיננטי?']},
    סגנון: {row['אתה בסגנון עירוני או ישובי?']}
    
    בזמני הפנוי- {row['מה משמח אותך? מה אתה אוהב לעשות? מה אתה עושה בזמן פנוי או חופשות?']}
    
    שאיפות אידיאלים ותכנונים לעתיד:
    {row['אידיאלים, שליחות מדבר אלייך?']}
    {row['שאיפות ותכנונים לעתיד']}
    
    רמה וסגנון תורני:
    {row['ספר לנו על הרמה והסגנון התורני שלך']}
    רוצה לבנות:
    {row['איזה סגנון בית אתה רוצה לבנות?']}    

    מה אתה מחפש:
    {row['תכונות שמהותיות וקריטיות לך בבחורה']}
    
    בהרחבה: {row['וקצת יותר בהרחבה..']}

    הכי חשוב לי-
    {row['לסיכום אם היית צריך לציין בנקודות מהם שלושת הדברים שהכי חשובים לך בבחורה, מה הם?']}
    
    מבחינה תורנית-
    {row['מה אתה מחפש אצל בחורה מבחינה תורנית?']}

    טווח גילאים רלוונטי: {row['טווח גילאים רלוונטי']}

    הערות נוספות-
    {row['הערות נוספות או מידע שיכול להועיל?']}

    האם חשובה לך העדה?: {row['האם חשובה לך העדה?']}
    האם יש משהו שאפסול עליו מבחינה חיצונית?: {row['האם יש משהו שאפסול עליו מבחינה חיצונית?']}
    
    מסקנות מפגישות קודמות:
    {row['אם יש מסקנות מפגישות קודמות מה נכון וטוב לי, ומה פחות,  כתוב לנו. זה יועיל בלדייק אותנו.']}
    
    משהו נוסף:
    {row['משהו נוסף שתרצה להוסיף בהקשר של מה אתה מחפש?']}

    מוכן לשלוח תמונה?
    {row['האם אתה מוכן לשלוח תמונה שלך לבחורה במקרה והיא מעוניינת בכך? ']}
    
    """ 
    # Use Gemini API to refine the text
    improved_text = improve_text_with_gemini(gemini_api_key, template_text)

    return improved_text

# Function to change font in a paragraph
def change_font(paragraph, font_name="Arial", font_size=12):
    # Loop through all runs in the paragraph
    for run in paragraph.runs:
        run.font.name = font_name  # Set the font
        run.font.size = Pt(font_size)  # Set the font size (in points)
        run.font.bold = False  # Optionally set the font to not bold
        run.font.italic = False  # Optionally set the font to not italic
        run.font.underline = False  # Optionally set the font to not underlined

# Function to add a paragraph with RTL alignment
# Function to add a paragraph with RTL alignment
def add_rtl_paragraph(doc, text, is_headline=False):
    # Add paragraph
    para = doc.add_paragraph(text)
    
    change_font(para, "arial", font_size=12)

    # If it's a headline, set the text to bold
#    if is_headline:
#        for run in para.runs:
#            run.font.bold = True

    # Set alignment to right
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Set the RTL direction (this requires setting the 'right-to-left' property)
    run = para.add_run()
    run._r.set('rtl', '1')  # This sets the text direction to right-to-left
    
    return para


# Function to sanitize the file name
def sanitize_filename(filename):
    # Replace invalid characters with underscores or remove them
    return re.sub(r'[\\/*?:"<>|]', '', filename)

# Authenticate Google Drive service
drive_service = authenticate_drive()

# Process new responses
new_ids = []
for index, row in responses.iterrows():
    response_id = row['Response ID']
    if response_id not in processed_ids:  # Process only unprocessed responses
        summary = generate_text(row)
        
        # Create a .docx file using python-docx
        doc = Document()

        # Process images in the "צרף בבקשה שתי תמונות שלך" column
        # Process images in the "צרף בבקשה שתי תמונות שלך" column
        if 'צרף בבקשה שתי תמונות שלך' in row and pd.notna(row['צרף בבקשה שתי תמונות שלך']):
            print("Processing images...")
            image_urls = row['צרף בבקשה שתי תמונות שלך']
            file_ids = extract_file_ids(image_urls)
            
            image_paths = []  # Collect all image paths
            for file_id in file_ids:
                image_path = f"image_{file_id}.jpg"
                print(f"preparing to download image_{file_id}.jpg")
                download_image_from_drive(file_id, drive_service, image_path)
                image_paths.append(image_path)

            # Add images side by side without a table
            add_images_side_by_side(doc, image_paths)

            # Clean up downloaded images
            for image_path in image_paths:
                os.remove(image_path)

        # List of headers that should be bold
        headlines = ["-", "\n-"]

        # Split the text into paragraphs
        paragraphs = summary.split(":\n")
#        print(paragraphs)
        # Loop through the paragraphs
        for paragraph in paragraphs:
            # Check if the paragraph is a headline
            for headline in headlines:
                if paragraph.startswith(headline):
                    # Make the headline bold
                    add_rtl_paragraph(doc, paragraph, is_headline=True)
                    print("bold")
                    break
            else:
                # Regular text (not a headline)
                add_rtl_paragraph(doc, paragraph, is_headline=False)
        
        # Save the document locally as a .docx file
        output_file = f"generated_{row['שם פרטי']}_{row['שם משפחה']}_{row['מספר הטלפון שלך']}.docx"
        safe_output_file = sanitize_filename(output_file)
        doc.save(safe_output_file)

        # Upload the file to Google Drive as a Google Doc
        file_id = upload_to_drive(safe_output_file, drive_service)
        print(f"Uploaded profile for {row['שם פרטי']} {row['שם משפחה']} to Google Drive with ID: {file_id}")
        
        # Optionally delete the local file after uploading
        os.remove(safe_output_file)
        print(f"Profile for {row['שם פרטי']}{row['שם משפחה']}{row['מספר הטלפון שלך']} saved to {output_file}")
        
        new_ids.append(response_id)

# Update the processed IDs file
if new_ids:
    with open(processed_ids_file, 'a') as file:
        file.write('\n'.join(new_ids) + '\n')
